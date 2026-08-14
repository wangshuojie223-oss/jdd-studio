"""POST /api/script 冒烟测试：docx 上传 → 海报方案（LLM 打桩）。"""
from __future__ import annotations

import io

from fastapi.testclient import TestClient


def _docx_bytes() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_paragraph("《测试剧》 英文名：TEST ECHO")
    doc.add_paragraph("杰克：男，35岁，刑警，冷峻")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_script_endpoint(monkeypatch):
    from app import main, promptgen

    async def fake_generate(script_text, model=None, refs=None):
        assert "杰克" in script_text  # 剧本内容真的被解析进来了
        assert refs is not None      # 参考图清单参数已接通
        return {"character": "TEST ECHO", "face_base": "悬疑", "schemes": [
            {"name": f"方向{i}（{'亮调' if i <= 6 else '暗调'}）", "prompt": f"p{i}", "note": ""}
            for i in range(1, 9)
        ], "model": "fake", "raw": ""}

    monkeypatch.setattr(promptgen, "generate_poster_schemes", fake_generate)
    monkeypatch.setattr(main.promptgen, "generate_poster_schemes", fake_generate)
    # v1.4.0 起 /api/script 会先抽角色表——打桩避免真实 LLM 调用和写 refs/roster.json
    from app import roster as roster_mod

    async def fake_roster(script_text, model=None, caller=None):
        return []
    monkeypatch.setattr(roster_mod, "extract_roster", fake_roster)
    monkeypatch.setattr(main.roster_mod, "extract_roster", fake_roster)
    monkeypatch.setattr(roster_mod, "save_roster", lambda r: None)   # 不落盘
    monkeypatch.setattr(main.roster_mod, "save_roster", lambda r: None)
    client = TestClient(main.app)

    r = client.post("/api/script", files={"file": ("剧本.docx", _docx_bytes())}, data={"model": "gemini-3.6-flash"})
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["character"] == "TEST ECHO"
    assert body["roster_count"] == 0
    assert len(body["schemes"]) == 8
    assert "亮调" in body["schemes"][0]["name"] and "暗调" in body["schemes"][7]["name"]

    # 非 docx 拒绝
    r2 = client.post("/api/script", files={"file": ("x.txt", b"hello")})
    assert r2.status_code == 400


def test_script_extracts_roster_and_fills_pending(tmp_path, monkeypatch):
    from app import store, main, config, promptgen, vision, roster as roster_mod
    monkeypatch.setattr(store, "_con", None)
    monkeypatch.setattr(config, "resolve_path", lambda p: tmp_path / p)
    monkeypatch.setattr(main.config, "resolve_path", lambda p: tmp_path / p)

    async def fake_generate(script_text, model=None, refs=None):
        return {"character": "TEST", "face_base": "", "schemes": [
            {"name": f"s{i}", "prompt": f"p{i}", "note": ""} for i in range(1, 9)],
            "model": "fake", "raw": ""}
    monkeypatch.setattr(main.promptgen, "generate_poster_schemes", fake_generate)

    async def fake_roster(script_text, model=None, caller=None):
        return [{"name": "Evelyn Hart", "gender": "Female", "age": 22,
                 "identity": "失明千金", "appearance": "金棕长发"}]
    monkeypatch.setattr(roster_mod, "extract_roster", fake_roster)
    monkeypatch.setattr(main.roster_mod, "extract_roster", fake_roster)

    async def fake_batch(image_paths, roster_list, caller=None):
        return [{"name": "Evelyn Hart", "confidence": 0.9, "reason": "x"} for _ in image_paths]
    monkeypatch.setattr(vision, "recognize_batch", fake_batch)
    monkeypatch.setattr(main.vision, "recognize_batch", fake_batch)

    client = TestClient(main.app)
    # 先传一张无名参考图 → pending
    import base64
    png = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==")
    r = client.post("/api/refs", files={"file": ("e.png", png)}, data={"name": ""})
    assert r.json()["pending"] is True

    r = client.post("/api/script", files={"file": ("剧本.docx", _docx_bytes())})
    assert r.status_code == 200, r.text
    assert r.json()["roster_count"] == 1
    # roster.json 已缓存、pending 已补识别
    assert roster_mod.load_roster()[0]["name"] == "Evelyn Hart"
    lst = client.get("/api/refs").json()["refs"]
    assert lst[0]["name"] == "Evelyn Hart" and "pending" not in lst[0]


def test_script_roster_failure_does_not_block(tmp_path, monkeypatch):
    from app import store, main, config, roster as roster_mod
    monkeypatch.setattr(store, "_con", None)
    monkeypatch.setattr(config, "resolve_path", lambda p: tmp_path / p)
    monkeypatch.setattr(main.config, "resolve_path", lambda p: tmp_path / p)

    async def fake_generate(script_text, model=None, refs=None):
        return {"character": "TEST", "face_base": "", "schemes": [
            {"name": f"s{i}", "prompt": f"p{i}", "note": ""} for i in range(1, 9)],
            "model": "fake", "raw": ""}
    monkeypatch.setattr(main.promptgen, "generate_poster_schemes", fake_generate)

    async def boom(script_text, model=None, caller=None):
        raise RuntimeError("抽取失败")
    monkeypatch.setattr(roster_mod, "extract_roster", boom)
    monkeypatch.setattr(main.roster_mod, "extract_roster", boom)

    client = TestClient(main.app)
    r = client.post("/api/script", files={"file": ("剧本.docx", _docx_bytes())})
    assert r.status_code == 200 and r.json()["roster_count"] == 0  # 海报照常出
