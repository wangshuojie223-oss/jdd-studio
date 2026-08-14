"""docx 剧本解析测试：段落+表格提取、超长截断、坏文件报错。"""
from __future__ import annotations

from pathlib import Path


def _make_docx(path: Path, paragraphs, table_rows=None):
    from docx import Document
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=2)
        for i, (a, b) in enumerate(table_rows):
            t.rows[i].cells[0].text = a
            t.rows[i].cells[1].text = b
    doc.save(str(path))


def test_extract_paragraphs_and_tables(tmp_path):
    from app import scriptdoc
    fp = tmp_path / "s.docx"
    _make_docx(fp, ["《Silent Echo》", "第一集 雨夜", "杰克：男，35岁，刑警"], table_rows=[("英文名", "SILENT ECHO")])
    text = scriptdoc.extract_text(fp)
    assert "雨夜" in text and "杰克" in text and "SILENT ECHO" in text


def test_extract_truncates(tmp_path):
    from app import scriptdoc
    fp = tmp_path / "big.docx"
    _make_docx(fp, ["字" * 30000])
    assert len(scriptdoc.extract_text(fp)) <= 20050  # 截断 + 省略标记


def test_extract_rejects_bad_file(tmp_path):
    from app import scriptdoc
    fp = tmp_path / "x.docx"
    fp.write_text("不是docx")
    try:
        scriptdoc.extract_text(fp)
        assert False, "应该抛 ScriptDocError"
    except scriptdoc.ScriptDocError:
        pass
