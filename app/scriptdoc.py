"""剧本 Word 文档解析：提取全文文本（段落 + 表格），超长截断。"""
from __future__ import annotations

from pathlib import Path

from docx import Document

MAX_CHARS = 20_000


class ScriptDocError(Exception):
    pass


def extract_text(path: Path) -> str:
    try:
        doc = Document(str(path))
    except Exception as e:
        raise ScriptDocError(f"无法读取 Word 文档：{e}")
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:  # 人物表/信息表常在表格里
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ScriptDocError("文档里没有读到文字内容")
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n……（剧本过长，已截断）"
    return text
